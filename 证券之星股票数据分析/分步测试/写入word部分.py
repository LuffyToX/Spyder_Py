# -*- coding:utf-8 -*-
import docx
import os

def main():
    """ 将分析结果存入 Word """

    # 新建 Word 文件
    doc = docx.Document()

    # 添加一级标题
    title1 = '股票分析结果'
    doc.add_heading(title1, 1)

    # 添加段落
    words = '使得'
    doc.add_paragraph(words)


    # 添加图片
    doc.add_picture('成交额.jpg', width=docx.shared.Inches(5), height=docx.shared.Cm(10))
    # 宽 5 英尺，高 10 cm

    # 保存文档
    doc.save('股票数据分析结果.docx')

if __name__ == "__main__":
    os.chdir(r'C:\Users\Dell\Desktop')
    main()
