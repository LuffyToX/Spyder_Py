# -*- coding:utf-8 -*-
""" 股票市场数据分析链条 """

import re
import time
import random
import urllib.request
import urllib.error
import docx
import os

class Crawl_data:
    """ 爬取数据 （数据来源：证券之星）"""

    def __init__(self, url, max_page):
        self.url = url
        self.page = int(max_page)


    def get_page_url(self):
        """ 获取相应页面的 url """
        url_page = self.url + str(self.page) + '.html'
        #print(self.url)
        #print(url_page)
        return url_page


    def get_page_ori_code(self):
        """ 获取网页源码 """

        ''' 伪装成浏览器 '''
        # 以字典的形式设置 headers
        headers = {"Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8",
                   "Accept-Language": "zh-CN,zh;q=0.9",
                   "User-Agent": "Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/70.0.3538.102 Safari/537.36",
                   "Connection": "keep-alive"}
        # 创建 bulid_opener 对象
        opener = urllib.request.build_opener()

        # 建立空列表，为了以指定格式存储头信息
        headall = list()

        # 通过 for 循环遍历字典，构造出指定格式的 headers 信息
        for key, value in headers.items():
            item = (key, value)
            headall.append(item)
        # 将指定格式的 headers 信息添加好
        # 该格式为 opener.addheaders = [("", ""), ("", ""), ("", ""), ...]
        opener.addheaders = headall
        # 将 opener 安装为全局
        urllib.request.install_opener(opener)

        ''' 打开网页 '''
        try:
            # 设置超时时间为 100s
            html = opener.open(self.get_page_url(), timeout=100)
            #print(type(html))
            return html
        except Exception as e:
            print("出现异常 ——>" + str(e))


    def get_shares_data(self):
        """ 获取股票数据 """
        shares_total = list()  # 所有页面的股票数据
        for page in range(self.page):

            ''' 异常检测 '''
            try:
                html = self.get_page_ori_code()
            except urllib.error.URLError as e:
                if hasattr(e, "code") and hasattr(e, "reason"):
                    print("Page {0} has a fault and it's code is: {1}\nThat means: {2}"\
                          .format(page+1, e.code, e.reason))
                elif hasattr(e, "reason"):
                    print("Page %d has a fault and it's reason is: %s"%(page+1, e.reason))
            except urllib.error.HTTPError as e:
                if hasattr(e, "code") and hasattr(e, "reason"):
                    print("Page {0} has a fault and it's code is: {1}\nThat means: {2}"\
                          .format(page+1, e.code, e.reason))
                elif hasattr(e, "reason"):
                    print("Page %d has a fault and it's reason is: %s" % (page+1, e.reason))

            ''' 获取页码 '''
            html = html.read().decode('gbk')
            print('Get page %d sucessfully!'%(page+1))

            # 匹配主体
            pat_body = re.compile('<tbody[\s\S]*</tbody>')
            body = re.findall(pat_body, str(html))

            # 某页的股票数据
            pat_shares = re.compile('>(.*?)<')
            shares_page = re.findall(pat_shares, body[0])

            shares_total.extend(shares_page)
            # 每抓一页随机休眠几秒，防止被服务器中断
            time.sleep(random.randrange(1, 4))

        ''' 删除空白字符 '''
        #print(shares_total) 测试用
        shares_total = [x.strip() for x in shares_total if x.strip()!='']
        #print(shares_total) 测试用
        return shares_total


def align_print(string, distance, alignment='left'):
    """ 中英混搭时对齐输出 """

    # 计算出 string 的 'gbk' 码长度
    length = len(string.encode('gbk'))
    if distance > length:
        space_to_fill = distance - length
    else:
        space_to_fill = 0

    if alignment == 'left':
        string = string + ' ' * space_to_fill

    elif alignment == 'right':
        string = ' ' * space_to_fill + string

    elif alignment == 'center':
        string = ' ' * (distance // 2) + string + ' ' * (distance - distance // 2)

    return string


import openpyxl
from openpyxl.utils import get_column_letter

class Save_to_Excel:
    def __init__(self, title, data):
        self.title = title
        self.data = data

        # 新建工作簿
        self.wb = openpyxl.Workbook()
        # 获取活跃工作表
        self.sheet = self.wb.active


    def add_title(self):
        """ 向 Excel 中写入标题 """

        ''' 为当前工作表更名 '''
        self.sheet.title = "股市"

        sheet_list = self.wb.sheetnames
        print("当前工作簿拥有的工作表为：%s" % sheet_list)

        ''' 写入标题 '''
        self.sheet.append(self.title)
        print("**** 标题写入成功 ****")


    def add_data(self):
        """ 向 Excel 中写入数据 """

        code = list()
        simple = list()
        latest_price = list()
        ups_and_downs = list()
        ups_and_downs_val = list()
        ups_in_5mins = list()
        volume = list()
        busi_volume = list()
        turn_rate = list()
        amplitude = list()
        volume_ratio = list()
        comparison = list()
        market_rate = list()


        for i in range(0, len(self.data), 13):
            code.append(self.data[i])
            simple.append(self.data[i+1])
            latest_price.append(self.data[i+2])
            ups_and_downs.append(self.data[i+3])
            ups_and_downs_val.append(self.data[i+4])
            ups_in_5mins.append(self.data[i+5])
            volume.append(self.data[i+6])
            busi_volume.append(self.data[i+7])
            turn_rate.append(self.data[i+8])
            amplitude.append(self.data[i+9])
            volume_ratio.append(self.data[i+10])
            comparison.append(self.data[i+11])
            market_rate.append(self.data[i+12])


        for num in range(13):
            col = get_column_letter(num + 1)  # 列号
            for row in range(len(self.data)//13):  # 行号
                col_and_row = col + str(row + 2)  # (要从第二行开始，否则会覆盖标题)
                if num == 0:
                    self.sheet[col_and_row] = code[row]
                elif num == 1:
                    self.sheet[col_and_row] = simple[row]
                elif num == 2:
                    self.sheet[col_and_row] = latest_price[row]
                elif num == 3:
                    self.sheet[col_and_row] = ups_and_downs[row]
                elif num == 4:
                    self.sheet[col_and_row] = ups_and_downs_val[row]
                elif num == 5:
                    self.sheet[col_and_row] = ups_in_5mins[row]
                elif num == 6:
                    self.sheet[col_and_row] = volume[row]
                elif num == 7:
                    self.sheet[col_and_row] = busi_volume[row]
                elif num == 8:
                    self.sheet[col_and_row] = turn_rate[row]
                elif num == 9:
                    self.sheet[col_and_row] = amplitude[row]
                elif num == 10:
                    self.sheet[col_and_row] = volume_ratio[row]
                elif num == 11:
                    self.sheet[col_and_row] = comparison[row]
                else:
                    self.sheet[col_and_row] = market_rate[row]

        print("**** 数据写入成功 ****")


    def set_width_and_height(self, width, height):
        """ 设置单元格行高和列宽 """

        # 列宽
        for num in range(13):
            col = get_column_letter(num + 1)
            self.sheet.column_dimensions[col].width = width
        # 行高
        self.sheet.row_dimensions[1].height = 20 # 标题行
        for row in range(len(self.data)//13): # 数据行
            self.sheet.row_dimensions[row+2].height = height


    def freeze_and_save(self):
        """ 冻结标题行以及保存 """

        # 如果要实现该函数，则需要写入标题和数据
        self.add_title()
        self.add_data()
        self.set_width_and_height(15, 15)

        self.sheet.freeze_panes = 'A2'

        self.wb.save('股市.xlsx')


import matplotlib.pyplot as plt
from multiprocessing import Process

class Data_analysis:
    """ 可视化的数据分析 """

    def __init__(self, data):
        self.data = data


    def pre_precess(self, num):
        """ 预处理 """

        """ 以股价进行排序 """

        ''' 将列表变为每个子列表有13个数据的嵌套列表(因为股票信息有13列数据) '''
        n = 13
        lst = [self.data[i:i + n] for i in range(0, len(self.data), n)]
        # print(data, '\n', len(data), sep='')

        ''' 列表去重 '''
        data = []
        for i in lst:
            if not i in data:
                data.append(i)

        ''' 冒泡法排序(由小到大) '''
        pat = re.compile(r"\d+\.?\d*") # 用于提取数字
        length = len(data)
        # 一共进行(length-1)轮排序
        for i in range(0, length - 1):
            for j in range(0, length - 1 - i):
                if float(re.findall(pat, data[j][num])[0]) > float(re.findall(pat, data[j+1][num])[0]):  # 注意元素是字符串
                    data[j], data[j + 1] = data[j + 1], data[j]
        #for value in data:
        #    print(value[num])

        return data


    def price(self):
        """ 找出股价前 10 的公司（蛮不错的嘛）"""

        data = self.pre_precess(2) # 股价位于第 3 列
        # 解决 plt 中文显示问题
        plt.rcParams['font.sans-serif'] = ['SimHei']
        plt.rcParams['axes.unicode_minus'] = False

        plt.figure()
        plt.title("最新股价前10位公司")
        plt.xlabel("企业")
        plt.ylabel("股价")

        x = []
        y = []
        for i in range(len(data)):
            x.append(data[i][1])
            y.append(data[i][2])
        x = x[-10:]
        y = y[-10:]

        #plt.plot(x, y) # 折线图
        plt.scatter(x, y, edgecolor='none', c='purple', s=100) # 散点图

        # plt.show() 后实际上已经创建了一个新的空白的图片（坐标轴）,
        # 这时候再 plt.savefig() 就会保存这个新生成的空白图片
        plt.savefig("股价.jpg", bbox_inches='tight')
        #plt.show()


    def ups_and_downs(self):
        """ 找出涨跌幅最高的 10 家公司（这也太惨了） """
        data = self.pre_precess(3) # 涨跌幅位于第 4 列
        plt.rcParams['font.sans-serif'] = ['SimHei']
        plt.rcParams['axes.unicode_minus'] = False

        plt.figure()
        plt.title("涨跌幅最高的 10 家公司")
        plt.xlabel("企业")
        plt.ylabel("涨跌幅")

        x = []
        y = []
        for i in range(len(data)):
            x.append(data[i][1])
            y.append(data[i][3])
        x = x[-10:]
        y = y[-10:]
        plt.scatter(x, y, edgecolor='none', c='purple', s=100)
        plt.savefig("涨跌幅.jpg", bbox_inches='tight')
        #plt.show()


    def volume(self):
        """ 找出成交量最高的 10 家公司（卖的还不错）"""
        data = self.pre_precess(6)  # 成交量位于第 7 列
        plt.rcParams['font.sans-serif'] = ['SimHei']
        plt.rcParams['axes.unicode_minus'] = False

        plt.figure()
        plt.title("成交量最高的 10 家公司")
        plt.xlabel("企业")
        plt.ylabel("成交量")

        x = []
        y = []
        for i in range(len(data)):
            x.append(data[i][1])
            y.append(data[i][6])
        x = x[-10:]
        y = y[-10:]
        plt.scatter(x, y, edgecolor='none', c='purple', s=100)
        plt.savefig("成交量.jpg", bbox_inches='tight')
        #plt.show()


    def busi_volume(self):
        """ 找出成交额最高的 10 家公司（这就很赚了） """
        data = self.pre_precess(7)  # 涨跌额位于第 8 列
        plt.rcParams['font.sans-serif'] = ['SimHei']
        plt.rcParams['axes.unicode_minus'] = False

        plt.figure()
        plt.title("成交额最高的 10 家公司")
        plt.xlabel("企业")
        plt.ylabel("成交额")

        x = []
        y = []
        for i in range(len(data)):
            x.append(data[i][1])
            y.append(data[i][7])
        x = x[-10:]
        y = y[-10:]
        plt.scatter(x, y, edgecolor='none', c='purple', s=100)
        plt.savefig("成交额.jpg", bbox_inches='tight')
        #plt.show()


    def cha_rate(self):
        """ 找出换手率最高的 10 家公司（看来经营的不咋地） """
        data = self.pre_precess(8)  # 换手率位于第 9 列
        plt.rcParams['font.sans-serif'] = ['SimHei']
        plt.rcParams['axes.unicode_minus'] = False

        plt.figure()
        plt.title("换手率最高的 10 家公司")
        plt.xlabel("企业")
        plt.ylabel("换手率")

        x = []
        y = []
        for i in range(len(data)):
            x.append(data[i][1])
            y.append(data[i][8])
        x = x[-10:]
        y = y[-10:]
        plt.scatter(x, y, edgecolor='none', c='purple', s=100)
        plt.savefig("换手率.jpg", bbox_inches='tight')
        #plt.show()


    def save_pic(self):
        """ 保存分析得到的图片 """
        pr = Process(target=self.price())
        up = Process(target=self.ups_and_downs())
        vo = Process(target=self.volume())
        bu = Process(target=self.busi_volume())
        ch = Process(target=self.cha_rate())

        pr.start()
        up.start()
        vo.start()
        bu.start()
        ch.start()

if __name__ == "__main__":
    os.chdir(r'C:\Users\Dell\Desktop\Program\learn_python\Py_homework\大作业_席天亮_信计1601_41603408\由程序生成的文件')
    url = 'http://quote.stockstar.com/stock/ranklist_a_3_1_' # 该url不能用于直接搜索
    max_page = input("你想爬取前多少页的信息？")
    data = Crawl_data(url, max_page)
    shares = data.get_shares_data() # 获取数据列表

    ''' 尽管我尝试很多方法试图对齐中英混搭输出，但效果都不尽人意，此种方法是输出效果最佳的 '''
    # 这里只打印前 5 列
    print(align_print('代码', 15), align_print('简称', 15),\
          align_print('最新价', 15), align_print('涨跌幅', 15),\
          align_print('涨跌额', 15), align_print('5分钟涨幅', 15))
    for i in range(0, len(shares), 13):
        print(align_print(shares[i], 15), align_print(shares[i+1], 15),\
              align_print(shares[i+2], 15), align_print(shares[i+3], 15),\
              align_print(shares[i+4], 15), align_print(shares[i+5], 15))

    title = ['代码', '简称', '最新价', '涨跌幅', '涨跌额', '5分钟涨幅', \
             '成交量(手)', '成交额(万元)', '换手率', '振幅', '量比', '委比', '市盈率']
    Excel = Save_to_Excel(title, shares)
    Excel.freeze_and_save() # 这里如果本地存在相应文件，会报错，删除后再执行

    data_analysis = Data_analysis(shares)
    ''' 这里搞了很久,开始时同时保存图片总会出现错乱，原因是没有调用 plt.figure '''

    data_analysis.save_pic()

    """ 将分析结果存入 Word """
    doc = docx.Document()
    doc.add_heading('股票分析结果', 1)

    doc.add_paragraph('股价')
    doc.add_picture('股价.jpg', width=docx.shared.Inches(5), height=docx.shared.Cm(10))
    doc.add_paragraph('成交量')
    doc.add_picture('成交量.jpg', width=docx.shared.Inches(5), height=docx.shared.Cm(10))
    doc.add_paragraph('成交额')
    doc.add_picture('成交额.jpg', width=docx.shared.Inches(5), height=docx.shared.Cm(10))
    doc.add_paragraph('涨跌幅')
    doc.add_picture('涨跌幅.jpg', width=docx.shared.Inches(5), height=docx.shared.Cm(10))
    doc.add_paragraph('换手率')
    doc.add_picture('换手率.jpg', width=docx.shared.Inches(5), height=docx.shared.Cm(10))

    doc.save('股票数据分析结果.docx')